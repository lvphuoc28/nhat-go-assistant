"""
TRO LY NHAT GO v5.0 — He thong hoi dap noi bo
Chay: python app.py
Truy cap: https://localhost:5000
          https://<IP-may-tinh>:5000 (may khac trong cung mang)
"""

import os, json, re, pickle, io
from flask import Flask, request, jsonify, render_template, Response, stream_with_context
import anthropic
try:
    import requests as _http
    REQUESTS_AVAILABLE = True
except ImportError:
    REQUESTS_AVAILABLE = False

# Thu import python-docx
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Thu import gTTS
try:
    from gtts import gTTS
    GTTS_AVAILABLE = True
except ImportError:
    GTTS_AVAILABLE = False

# BM25
try:
    from rank_bm25 import BM25Okapi
    BM25_AVAILABLE = True
except ImportError:
    BM25_AVAILABLE = False

app = Flask(__name__)

# ── CAU HINH ─────────────────────────────────────────────────────────────────
BASE_DIR     = os.path.dirname(os.path.abspath(__file__))
CONFIG_FILE  = os.path.join(BASE_DIR, "config.json")
KB_FILE      = os.path.join(BASE_DIR, "SoQuyDinhTongHop_NhatGo.md")
DOCX_KB_FILE = os.path.join(BASE_DIR, "ToanBoQuyDinh_NhatGo_2026.docx")
INDEX_CACHE  = os.path.join(BASE_DIR, "bm25_index.pkl")
FOLDER7_PATH = os.path.join(BASE_DIR, "..", "..", "1. MAU-NS", "7.QUI DINH CTY NHAT GO")
SSL_CERT     = os.path.join(BASE_DIR, "server.crt")
SSL_KEY      = os.path.join(BASE_DIR, "server.key")

def load_config():
    cfg = {}
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            try:
                cfg = json.load(f)
            except Exception:
                cfg = {}
    # Doc tu bien moi truong (Render.com / cloud deployment)
    env_map = {
        'api_key':            'ANTHROPIC_API_KEY',
        'zalo_app_id':        'ZALO_APP_ID',
        'zalo_app_secret':    'ZALO_APP_SECRET',
        'zalo_access_token':  'ZALO_ACCESS_TOKEN',
        'zalo_refresh_token': 'ZALO_REFRESH_TOKEN',
        'zalo_oa_id':         'ZALO_OA_ID',
    }
    for cfg_key, env_key in env_map.items():
        val = os.environ.get(env_key, '')
        if val:
            cfg[cfg_key] = val
    return cfg

def save_config(data):
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# ── TU DONG NGHIA ────────────────────────────────────────────────────────────
SYNONYMS = {
    'xin con': 'sinh con', 'vo de': 'thai san', 'de con': 'thai san',
    'vo sinh': 'thai san', 'sinh em be': 'thai san', 'bau': 'thai san',
    'mat': 'qua doi', 'chet': 'qua doi', 'ba mat': 'cha mat',
    'ma mat': 'me mat', 'thoi viec': 'nghi viec', 'bi duoi': 'sa thai',
    'bi sa thai': 'ky luat', 'tien thuong': 'thuong', 'benh': 'om',
    'vao vien': 'nam vien', 'di benh vien': 'nam vien',
    'di muon': 'di tre', 'den muon': 'di tre', 'muon': 'tre',
    'len bac': 'nang bac', 'tang bac': 'nang bac',
    'xin ve': 'giay ra cong', 'ra ve som': 'giay ra cong', 've som': 'giay ra cong',
}

def normalize_query(text):
    t = text.lower()
    for slang, official in SYNONYMS.items():
        t = t.replace(slang, official)
    return t

STOPWORDS_BM25 = {
    'toi','ban','anh','chi','em','minh','ho','chung','no',
    'va','hay','hoac','nhung','ma','vi','nen','thi','la','cua',
    'co','khong','duoc','cho','voi','ve','trong','tren','duoi',
    'khi','neu','sao','gi','nao','day','do','nay','kia',
    'bao','nhieu','it','lam','qua','rat','cung','deu','da',
    'dang','se','van','con','toi','den','nhu','theo','sau',
    'truoc','luc','ngay','luon','moi','cac','nhung','mot','hai',
    'lam','noi','hoi','xin','muon','can','phai','biet',
    'oi','the','vay','nha','nhe','bi','do','tu','ra','vo',
    'ty','cong','dieu','khoan','lao',
}

def tokenize_vi(text):
    text = text.lower()
    text = re.sub(r'[^\w\s]', ' ', text)
    tokens = text.split()
    tokens = [t for t in tokens if t not in STOPWORDS_BM25 and len(t) >= 2]
    return tokens if tokens else text.split()

# ── DOC TAI LIEU ─────────────────────────────────────────────────────────────
def extract_chunks_from_md(filepath):
    chunks = []
    if not os.path.exists(filepath):
        return chunks
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    sections = re.split(r'#{6,}', content)
    for sec in sections:
        s = sec.strip()
        if len(s) < 200:
            continue
        lines = [l.strip() for l in s.split('\n') if l.strip()]
        if len(lines) <= 3 and any('gom' in l.lower() for l in lines):
            continue
        title = next((l for l in lines if len(l) > 3), 'Quy dinh')
        title = re.sub(r'^#+\s*', '', title).strip('- ')
        chunks.append({
            'title': title[:120], 'text': s,
            'source': 'So Quy Dinh Tong Hop',
            'year': '2026', 'tokens': tokenize_vi(s),
        })
    return chunks

def extract_chunks_from_docx(filepath):
    chunks = []
    if not DOCX_AVAILABLE or not os.path.exists(filepath):
        return chunks
    try:
        doc = DocxDocument(filepath)
        current_title, current_meta, current_body = '', '', []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name if para.style else ''
            if style.startswith('Heading 2'):
                if current_body and len('\n'.join(current_body)) > 150:
                    full = current_title + '\n' + current_meta + '\n' + '\n'.join(current_body)
                    chunks.append({
                        'title': current_title[:120], 'text': full,
                        'source': current_title[:80],
                        'year': _extract_year(current_meta),
                        'tokens': tokenize_vi(full),
                    })
                current_title, current_meta, current_body = text, '', []
            elif style.startswith('Heading 1'):
                current_title, current_body = '', []
            else:
                if 'Nam ban hanh' in text or 'File goc' in text:
                    current_meta = text
                elif text != '[File .doc cu - can mo truc tiep de xem noi dung]':
                    current_body.append(text)
        if current_body and len('\n'.join(current_body)) > 150:
            full = current_title + '\n' + current_meta + '\n' + '\n'.join(current_body)
            chunks.append({
                'title': current_title[:120], 'text': full,
                'source': current_title[:80],
                'year': _extract_year(current_meta),
                'tokens': tokenize_vi(full),
            })
    except Exception as e:
        print(f"[DOCX] Loi doc: {e}")
    return chunks

def _extract_year(meta_line):
    m = re.search(r'(201[89]|202\d)', meta_line)
    return m.group(1) if m else '?'

_SUBFOLDER_YEAR_MAP = {
    '11. Qui dinh 2026': '2026',
    '7. Qui dinh 2025': '2025',
    '6. Qui dinh 2024': '2024',
    '5. Qui dinh 2023': '2023',
    '4. qui dinh 2022': '2022',
    '2.Qui dinh 2021': '2021',
    '2. Qui dinh 2021': '2021',
    '3. Quy dinh nam 2020': '2020',
    '1. Qui dinh Nam 2019': '2019',
    '8. Qui dinh ve luong': '?',
    '9. QUY Y TE': '?',
    '10. Tong hop qui dinh': '?',
    'QUY DINH DOI VOI NHA AN': '?',
    'THOI VU': '?',
}

def extract_year_from_path(filepath):
    norm = filepath.replace('\\', '/')
    for folder, year in _SUBFOLDER_YEAR_MAP.items():
        if folder.lower() in norm.lower():
            return year
    fname = os.path.basename(norm)
    m = re.search(r'(201[6-9]|202[0-6])', fname)
    if m:
        return m.group(1)
    m = re.search(r'(201[6-9]|202[0-6])', norm)
    if m:
        return m.group(1)
    return '2018'

def load_folder7_all_docs(folder7_path):
    chunks = []
    if not DOCX_AVAILABLE or not os.path.exists(folder7_path):
        print(f"[FOLDER7] Khong tim thay: {folder7_path}")
        return chunks
    ok, skip = 0, 0
    for root, dirs, files in os.walk(folder7_path):
        dirs[:] = [d for d in dirs if not d.startswith('~$') and d != '12. SCAN QUY DINH 2026']
        for fname in sorted(files):
            if fname.startswith('~$') or fname.lower() == 'thumbs.db':
                continue
            if not (fname.lower().endswith('.docx') or fname.lower().endswith('.doc')):
                continue
            fpath = os.path.join(root, fname)
            year = extract_year_from_path(fpath)
            try:
                doc = DocxDocument(fpath)
                paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                if len(paras) < 2:
                    skip += 1
                    continue
                full_text = '\n'.join(paras)
                if len(full_text) < 80:
                    skip += 1
                    continue
                raw_title = fname.replace('.docx', '').replace('.doc', '')
                title = re.sub(r'^[\d\-\.]+\s*', '', raw_title).strip() or raw_title
                title = title[:100]
                source_label = f"Quy dinh {year} -- {title}" if year != '?' else title
                chunks.append({
                    'title': title, 'text': full_text,
                    'source': source_label, 'year': year,
                    'tokens': tokenize_vi(full_text),
                })
                ok += 1
            except Exception:
                skip += 1
    print(f"[FOLDER7] Xong: {ok} van ban hop le, bo qua {skip} file")
    return chunks

# ── XAY DUNG INDEX ────────────────────────────────────────────────────────────
def build_index():
    if os.path.exists(INDEX_CACHE):
        try:
            with open(INDEX_CACHE, 'rb') as f:
                data = pickle.load(f)
            if data.get('version') == '4.0' and len(data.get('chunks', [])) > 0:
                print(f"[INDEX] Da load cache: {len(data['chunks'])} chunks")
                return data['bm25'], data['chunks']
        except Exception:
            pass

    print("[INDEX] Dang xay dung BM25 index tu ToanBoQuyDinh_NhatGo_2026.docx...")
    chunks = []
    docx_chunks = extract_chunks_from_docx(DOCX_KB_FILE)
    chunks.extend(docx_chunks)
    print(f"[INDEX] DOCX: {len(docx_chunks)} chunks")

    if not chunks:
        print("[INDEX] CANH BAO: Khong co chunk nao!")
        return None, []

    tokenized = [c['tokens'] for c in chunks]
    bm25 = BM25Okapi(tokenized)
    print(f"[INDEX] BM25 san sang: {len(chunks)} van ban (2018-2026)")

    try:
        with open(INDEX_CACHE, 'wb') as f:
            pickle.dump({'version': '4.0', 'bm25': bm25, 'chunks': chunks}, f)
    except Exception as e:
        print(f"[INDEX] Khong luu cache duoc: {e}")

    return bm25, chunks

def _year_sort_key(year_str):
    try:
        return int(year_str)
    except (ValueError, TypeError):
        return -1

def search(question, bm25, chunks, top_k=8, max_chars=8000):
    if bm25 is None or not chunks:
        return "", []
    q_normalized = normalize_query(question)
    query_tokens = tokenize_vi(q_normalized) or tokenize_vi(question)
    scores = bm25.get_scores(query_tokens)
    top_idx = sorted(range(len(scores)), key=lambda i: -scores[i])[:top_k * 3]
    candidates = [(chunks[i], scores[i]) for i in top_idx if scores[i] > 0.01]
    if not candidates:
        return "", []
    candidates.sort(key=lambda x: (-_year_sort_key(x[0]['year']), -x[1]))
    results = candidates[:top_k]
    context = ""
    for chunk, score in results:
        snippet = chunk['text'][:2000]
        year_label = chunk['year'] if chunk['year'] != '?' else 'truoc 2019'
        context += f"\n\n[NAM {year_label} | {chunk['title']}]\n{snippet}"
        if len(context) > max_chars:
            break
    return context.strip(), results

# ── KHOI DONG INDEX ───────────────────────────────────────────────────────────
if BM25_AVAILABLE:
    BM25_INDEX, CHUNKS = build_index()
else:
    print("[INDEX] rank_bm25 chua cai. Chay: pip install rank_bm25")
    BM25_INDEX, CHUNKS = None, []

# ── HTTPS TU KY ───────────────────────────────────────────────────────────────
def ensure_ssl():
    """Tao chung chi HTTPS tu ky neu chua co."""
    if os.path.exists(SSL_CERT) and os.path.exists(SSL_KEY):
        return True
    try:
        from cryptography import x509
        from cryptography.x509.oid import NameOID
        from cryptography.hazmat.primitives import hashes, serialization
        from cryptography.hazmat.primitives.asymmetric import rsa
        import datetime, ipaddress, socket

        hostname = socket.gethostname()
        try:
            local_ip = socket.gethostbyname(hostname)
        except Exception:
            local_ip = '127.0.0.1'

        key = rsa.generate_private_key(public_exponent=65537, key_size=2048)
        subject = issuer = x509.Name([
            x509.NameAttribute(NameOID.COMMON_NAME, u"NhatGo-Local"),
        ])
        san_list = [
            x509.DNSName(u"localhost"),
            x509.IPAddress(ipaddress.IPv4Address(u"127.0.0.1")),
        ]
        try:
            san_list.append(x509.IPAddress(ipaddress.IPv4Address(local_ip)))
        except Exception:
            pass

        cert = (x509.CertificateBuilder()
            .subject_name(subject)
            .issuer_name(issuer)
            .public_key(key.public_key())
            .serial_number(x509.random_serial_number())
            .not_valid_before(datetime.datetime.utcnow())
            .not_valid_after(datetime.datetime.utcnow() + datetime.timedelta(days=3650))
            .add_extension(x509.SubjectAlternativeName(san_list), critical=False)
            .sign(key, hashes.SHA256()))

        with open(SSL_KEY, 'wb') as f:
            f.write(key.private_bytes(
                serialization.Encoding.PEM,
                serialization.PrivateFormat.TraditionalOpenSSL,
                serialization.NoEncryption()
            ))
        with open(SSL_CERT, 'wb') as f:
            f.write(cert.public_bytes(serialization.Encoding.PEM))

        print(f"[SSL] Da tao chung chi HTTPS cho {local_ip}")
        return True
    except ImportError:
        print("[SSL] Chua co thu vien 'cryptography' -- chay HTTP thay the")
        return False
    except Exception as e:
        print(f"[SSL] Loi tao cert: {e} -- chay HTTP thay the")
        return False

# ── ROUTES ────────────────────────────────────────────────────────────────────
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/status')
def status():
    cfg = load_config()
    return jsonify({"configured": bool(cfg.get("api_key", "").strip())})

@app.route('/api/get-key')
def get_key():
    cfg = load_config()
    key = cfg.get("api_key", "")
    masked = (key[:7] + "..." + key[-4:]) if len(key) > 8 else key
    return jsonify({"key": masked})

@app.route('/api/save-key', methods=['POST'])
def save_key():
    data = request.get_json()
    key = data.get("api_key", "").strip()
    cfg = load_config()
    cfg["api_key"] = key
    save_config(cfg)
    return jsonify({"ok": True})

@app.route('/api/tts', methods=['POST'])
def tts():
    if not GTTS_AVAILABLE:
        return jsonify({"error": "gTTS chua cai. Chay: pip install gtts"}), 500
    data = request.get_json()
    text = (data.get('text', '') or '').strip()
    if not text:
        return jsonify({"error": "Khong co text"}), 400
    try:
        tts_obj = gTTS(text=text, lang='vi', slow=False)
        buf = io.BytesIO()
        tts_obj.write_to_fp(buf)
        buf.seek(0)
        return Response(buf.read(), content_type='audio/mpeg',
                        headers={'Cache-Control': 'no-cache'})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/chat', methods=['POST'])
def chat():
    cfg = load_config()
    api_key = cfg.get("api_key", "").strip()
    if not api_key:
        return jsonify({"error": "Chua cai API Key. Nhan nut cai dat de thiet lap."}), 400

    data = request.get_json()
    question = data.get("question", "").strip()
    if not question:
        return jsonify({"error": "Cau hoi trong."}), 400

    context, results = search(question, BM25_INDEX, CHUNKS)
    import json as _json

    if not context:
        def _no_ctx():
            msg = "Cau hoi nay toi chua tim thay thong tin trong tai lieu noi bo. Anh/chi/em vui long lien he phong Nhan Su de duoc ho tro nhe."
            yield 'data: ' + _json.dumps({"t": "txt", "d": msg}) + '\n\n'
            yield 'data: ' + _json.dumps({"t": "done", "src": ""}) + '\n\n'
        return Response(stream_with_context(_no_ctx()),
                        content_type='text/event-stream',
                        headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})

    source_label = results[0][0]['title'][:60] if results else "So Quy Dinh Nhat Go"

    system_prompt = (
        "Ban la Tro Ly Nhat Go -- tra loi cau hoi nhan vien Cong ty TNHH MTV Nhat Go.\n\n"
        "QUY TAC BAT BUOC:\n"
        "1. CHI dung thong tin trong VAN BAN QUY DINH ben duoi. KHONG dung kien thuc ben ngoai.\n"
        "2. Van ban sap xep MOI NHAT -> CU NHAT. Quy dinh nam moi hon co HIEU LUC CAO HON.\n"
        "3. Neu co nhieu nam: tra loi theo quy dinh MOI NHAT truoc, neu ro su thay doi qua cac nam neu co.\n"
        "4. Neu day du: nam ban hanh, so ngay, so lan, muc tien, dieu kien cu the.\n"
        "5. Neu khong co trong tai lieu: noi 'Cau hoi nay chua co thong tin trong tai lieu, "
        "anh/chi/em lien he Nhan Su nhe.'\n"
        "6. Cuoi cau tra loi ghi: (Nguon: ten van ban, nam)\n\n"
        "DINH DANG -- RAT QUAN TRONG:\n"
        "TUYET DOI KHONG dung markdown: khong dau gach ngang (-), khong dau sao (*), "
        "khong ## tieu de, khong **in dam**.\n"
        "Liet ke: dung so thu tu '1. 2. 3.' hoac dau cham phay (;).\n"
        "Cau phai co DAY DU chu ngu va vi ngu. "
        "VI DU SAI: 'Nghi 1 lan/thang.' -- DUNG: 'Anh/chi/em duoc nghi 1 lan moi thang.'\n"
        "Dung tu noi tu nhien: 'Cu the la', 'Ngoai ra', 'Theo quy dinh', 'Doi voi truong hop'.\n"
        "Giong am ap, ro rang nhu HR dang giai thich truc tiep.\n"
        "Xung 'toi', goi nhan vien la 'anh/chi/em'.\n"
        "Tra loi DAY DU -- KHONG cat ngan, KHONG bo sot dieu kien quan trong.\n\n"
        "VAN BAN QUY DINH NHAT GO (sap xep moi -> cu):\n"
        + context
    )

    def generate():
        try:
            client = anthropic.Anthropic(api_key=api_key)
            with client.messages.stream(
                model="claude-haiku-4-5-20251001",
                max_tokens=1500,
                system=system_prompt,
                messages=[{"role": "user", "content": question}]
            ) as stream:
                for txt in stream.text_stream:
                    yield 'data: ' + _json.dumps({"t": "txt", "d": txt}) + '\n\n'
            yield 'data: ' + _json.dumps({"t": "done", "src": source_label}) + '\n\n'
        except anthropic.AuthenticationError:
            yield 'data: ' + _json.dumps({"t": "err", "d": "API Key khong hop le. Kiem tra lai trong Cai dat."}) + '\n\n'
        except Exception as e:
            yield 'data: ' + _json.dumps({"t": "err", "d": str(e)}) + '\n\n'

    return Response(stream_with_context(generate()),
                    content_type='text/event-stream',
                    headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})

# ── ZALO OA INTEGRATION ──────────────────────────────────────────────────────
def zalo_send(user_id, text):
    """Gui tin nhan den user qua Zalo OA API."""
    if not REQUESTS_AVAILABLE:
        print("[ZALO] requests chua cai")
        return False
    cfg = load_config()
    token = cfg.get('zalo_access_token', '')
    if not token:
        print("[ZALO] Chua co access token")
        return False
    if len(text) > 2000:
        text = text[:1997] + '...'

    def _post(tok):
        return _http.post(
            'https://openapi.zalo.me/v3.0/oa/message/cs',
            headers={'access_token': tok, 'Content-Type': 'application/json'},
            json={'recipient': {'user_id': user_id}, 'message': {'text': text}},
            timeout=10
        ).json()

    res = _post(token)
    # Token het han -> thu refresh
    if res.get('error') in (-216, -124, 10, -13):
        new_token = zalo_refresh()
        if new_token:
            res = _post(new_token)
    ok = (res.get('error') == 0)
    if not ok:
        print(f"[ZALO] Loi gui: {res}")
    return ok

# Luu cau tra loi cuoi cung cho moi user (de tao audio khi can)
_last_answers = {}  # {user_id: answer_text}

# ── LOG CAU HOI ───────────────────────────────────────────────────────────────
import datetime as _dt
_qa_log = []   # [{time, user_id, question, answer}]
LOG_FILE = os.path.join(BASE_DIR, "qa_log.json")

def log_qa(user_id, question, answer):
    """Ghi log cau hoi va cau tra loi."""
    entry = {
        'time': _dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'user_id': user_id,
        'question': question,
        'answer': answer,
    }
    _qa_log.append(entry)
    # Luu xuong file (giu lai khi restart neu filesystem con)
    try:
        with open(LOG_FILE, 'w', encoding='utf-8') as f:
            json.dump(_qa_log, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def load_log_from_file():
    """Doc log tu file neu co (sau khi restart)."""
    if os.path.exists(LOG_FILE):
        try:
            with open(LOG_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                _qa_log.extend(data)
                print(f"[LOG] Da doc {len(data)} ban ghi cu")
        except Exception:
            pass

load_log_from_file()

def zalo_generate_audio(text):
    """Tao file mp3 tu text bang gTTS."""
    if not GTTS_AVAILABLE:
        return None
    try:
        clean = text[:600]
        tts = gTTS(text=clean, lang='vi', slow=False)
        buf = io.BytesIO()
        tts.write_to_fp(buf)
        buf.seek(0)
        return buf.read()
    except Exception as e:
        print(f"[TTS] Loi tao audio: {e}")
        return None

def zalo_upload_audio(audio_bytes, token):
    """Upload file mp3 len Zalo, tra ve file_id."""
    try:
        res = _http.post(
            'https://upload.zalo.me/v2.0/oa/upload/file',
            headers={'access_token': token},
            files={'file': ('traLoi.mp3', audio_bytes, 'audio/mpeg')},
            timeout=30
        ).json()
        file_id = (res.get('data') or {}).get('file_id', '')
        if file_id:
            print(f"[ZALO] Upload audio OK: {file_id[:20]}...")
        else:
            print(f"[ZALO] Upload audio loi: {res}")
        return file_id
    except Exception as e:
        print(f"[ZALO] Loi upload audio: {e}")
        return ''

def zalo_send_audio(user_id, text):
    """Tao mp3 tu text va gui cho user qua Zalo."""
    audio_bytes = zalo_generate_audio(text)
    if not audio_bytes:
        zalo_send(user_id, 'Xin loi, chuc nang nghe hien chua kha dung.')
        return False
    cfg = load_config()
    token = cfg.get('zalo_access_token', '')
    file_id = zalo_upload_audio(audio_bytes, token)
    if not file_id:
        zalo_send(user_id, 'Xin loi, khong the tao file audio. Vui long doc text phia tren.')
        return False
    try:
        res = _http.post(
            'https://openapi.zalo.me/v3.0/oa/message/cs',
            headers={'access_token': token, 'Content-Type': 'application/json'},
            json={
                'recipient': {'user_id': user_id},
                'message': {'attachment': {'type': 'file', 'payload': {'file_id': file_id}}}
            },
            timeout=20
        ).json()
        if res.get('error') == 0:
            print(f"[ZALO] Da gui audio cho {user_id}")
            return True
        print(f"[ZALO] Loi gui audio: {res}")
        return False
    except Exception as e:
        print(f"[ZALO] Loi gui audio: {e}")
        return False

def zalo_refresh():
    """Lam moi OA Access Token bang Refresh Token."""
    if not REQUESTS_AVAILABLE:
        return None
    cfg = load_config()
    r_token  = cfg.get('zalo_refresh_token', '')
    app_id   = cfg.get('zalo_app_id', '')
    secret   = cfg.get('zalo_app_secret', '')
    if not all([r_token, app_id, secret]):
        return None
    try:
        res = _http.post(
            'https://oauth.zaloapp.com/v4/oa/access_token',
            data={'refresh_token': r_token, 'app_id': app_id, 'grant_type': 'refresh_token'},
            headers={'secret_key': secret},
            timeout=10
        ).json()
        new_tok = res.get('access_token', '')
        if new_tok:
            cfg['zalo_access_token'] = new_tok
            if res.get('refresh_token'):
                cfg['zalo_refresh_token'] = res['refresh_token']
            save_config(cfg)
            # Quan trong: cap nhat os.environ de load_config() dung token moi
            os.environ['ZALO_ACCESS_TOKEN'] = new_tok
            if res.get('refresh_token'):
                os.environ['ZALO_REFRESH_TOKEN'] = res['refresh_token']
            print(f"[ZALO] Da lam moi access token: {new_tok[:20]}...")
            return new_tok
    except Exception as e:
        print(f"[ZALO] Loi refresh token: {e}")
    return None

_ZALO_SYSTEM = (
    "Ban la Tro Ly Nhat Go -- tra loi cau hoi nhan vien Cong ty TNHH MTV Nhat Go.\n\n"
    "QUY TAC BAT BUOC:\n"
    "1. CHI dung thong tin trong VAN BAN QUY DINH ben duoi. KHONG dung kien thuc ben ngoai.\n"
    "2. Van ban sap xep MOI NHAT -> CU NHAT. Quy dinh nam moi hon co HIEU LUC CAO HON.\n"
    "3. Neu day du: nam ban hanh, so ngay, so lan, muc tien, dieu kien cu the.\n"
    "4. Neu khong co: noi 'Cau hoi nay chua co thong tin, anh/chi/em lien he Nhan Su nhe.'\n"
    "5. Cuoi tra loi ghi: (Nguon: ten van ban, nam)\n\n"
    "DINH DANG:\n"
    "KHONG dung markdown. Liet ke bang so thu tu. Gion am ap nhu HR giai thich truc tiep.\n"
    "Xung 'toi', goi nhan vien la 'anh/chi/em'.\n"
    "Tra loi NGAN GON phu hop voi chat (toi da 400 tu).\n\n"
    "VAN BAN QUY DINH:\n"
)

@app.route('/zalo_verifierEVwu6RxHR2L3s-mRjgyuH5A9iqo-_sOLCZam.html')
def zalo_verify_domain():
    return 'zalo-platform-site-verification: EVwu6RxHR2L3s-mRjgyuH5A9iqo-_sOLCZam'

@app.route('/zalo/webhook', methods=['GET', 'POST'])
def zalo_webhook():
    if request.method == 'GET':
        # Xac thuc webhook URL voi Zalo
        return jsonify({'status': 'ok'})

    data = request.get_json(silent=True) or {}
    event = data.get('event_name', '')
    print(f"[ZALO] Event: {event}")

    # Chi xu ly tin nhan text tu user
    if event != 'user_send_text':
        return jsonify({'status': 'ok'})

    sender_id = (data.get('sender') or {}).get('id', '')
    msg_text  = (data.get('message') or {}).get('text', '').strip()

    if not sender_id or not msg_text:
        return jsonify({'status': 'ok'})

    print(f"[ZALO] Hoi: {msg_text[:80]}")

    # Xu ly lenh nghe audio
    _NGHE_KEYWORDS = {'nghe', '🔊', 'nghe audio', 'doc cho toi nghe', 'phat audio'}
    if msg_text.lower().strip() in _NGHE_KEYWORDS:
        last = _last_answers.get(sender_id, '')
        if last:
            zalo_send(sender_id, '🔊 Dang tao audio, vui long cho giay lat...')
            zalo_send_audio(sender_id, last)
        else:
            zalo_send(sender_id, 'Chua co cau tra loi nao de phat. Anh/chi/em hay hoi cau hoi truoc nhe.')
        return jsonify({'status': 'ok'})

    cfg = load_config()
    api_key = cfg.get('api_key', '').strip()
    if not api_key:
        zalo_send(sender_id, 'He thong dang bao tri. Vui long lien he Nhan Su.')
        return jsonify({'status': 'ok'})

    context, results = search(msg_text, BM25_INDEX, CHUNKS)
    if not context:
        zalo_send(sender_id,
            'Cau hoi nay toi chua tim thay thong tin trong tai lieu noi bo. '
            'Anh/chi/em vui long lien he phong Nhan Su de duoc ho tro nhe.')
        return jsonify({'status': 'ok'})

    source = results[0][0]['title'][:60] if results else ''
    try:
        client = anthropic.Anthropic(api_key=api_key)
        resp = client.messages.create(
            model='claude-haiku-4-5-20251001',
            max_tokens=800,
            system=_ZALO_SYSTEM + context,
            messages=[{'role': 'user', 'content': msg_text}]
        )
        answer = resp.content[0].text
        log_qa(sender_id, msg_text, resp.content[0].text)
        if source:
            answer += f'\n\n(Nguon: {source})'
        zalo_send(sender_id, answer)
        print(f"[ZALO] Da tra loi ({len(answer)} ky tu)")
    except Exception as e:
        print(f"[ZALO] Loi Claude: {e}")
        zalo_send(sender_id, 'Co loi xay ra. Vui long thu lai sau.')

    return jsonify({'status': 'ok'})

@app.route('/zalo/oauth')
def zalo_oauth():
    """Nhan code tu Zalo OAuth va doi lay access token moi."""
    code = request.args.get('code', '')
    if not code:
        return '<h2>Khong co code. Vui long thu lai.</h2>', 400
    cfg = load_config()
    try:
        res = _http.post(
            'https://oauth.zaloapp.com/v4/oa/access_token',
            data={
                'code': code,
                'app_id': cfg.get('zalo_app_id', ''),
                'grant_type': 'authorization_code',
            },
            headers={'secret_key': cfg.get('zalo_app_secret', '')},
            timeout=10
        ).json()
        if res.get('access_token'):
            cfg['zalo_access_token'] = res['access_token']
            if res.get('refresh_token'):
                cfg['zalo_refresh_token'] = res['refresh_token']
            save_config(cfg)
            print(f"[ZALO] Da cap nhat token moi qua OAuth")
            return f'''<h2>&#10003; Token da luu thanh cong!</h2>
<p>Access token: {res["access_token"][:30]}...</p>
<p>Hay copy access_token va refresh_token vao Render.com Environment Variables:</p>
<pre>ZALO_ACCESS_TOKEN={res["access_token"]}</pre>
<pre>ZALO_REFRESH_TOKEN={res.get("refresh_token","")}</pre>
<p><a href="/">Quay ve trang chu</a></p>'''
        return f'<h2>Loi: {res}</h2>', 400
    except Exception as e:
        return f'<h2>Loi: {e}</h2>', 500

@app.route('/api/save-zalo', methods=['POST'])
def save_zalo_cfg():
    data = request.get_json()
    cfg = load_config()
    cfg['zalo_access_token'] = data.get('access_token', '')
    cfg['zalo_refresh_token'] = data.get('refresh_token', '')
    cfg['zalo_app_id']        = data.get('app_id', '')
    cfg['zalo_app_secret']    = data.get('app_secret', '')
    save_config(cfg)
    return jsonify({'ok': True})

# ── TRANG LOG CAU HOI ────────────────────────────────────────────────────────
LOG_PASSWORD = os.environ.get('LOG_PASSWORD', 'nhatgo2026')

@app.route('/logs')
def logs_page():
    pw = request.args.get('pw', '')
    if pw != LOG_PASSWORD:
        return '''<html><body style="font-family:Arial;padding:30px">
<h2>Xem log cau hoi Zalo</h2>
<form><input name="pw" type="password" placeholder="Mat khau" style="padding:8px;font-size:16px">
<button type="submit" style="padding:8px 16px;font-size:16px">Dang nhap</button></form>
</body></html>'''
    total = len(_qa_log)
    rows = ''
    for e in reversed(_qa_log[-100:]):
        q = e['question'][:80].replace('<','&lt;')
        a = e['answer'][:120].replace('<','&lt;')
        rows += f"<tr><td>{e['time']}</td><td>{q}</td><td>{a}...</td></tr>"
    return f'''<html><body style="font-family:Arial;padding:20px">
<h2>Log cau hoi Zalo — Nhat Go ({total} ban ghi)</h2>
<p><a href="/logs/download?pw={pw}" style="background:#4CAF50;color:white;padding:10px 20px;
text-decoration:none;border-radius:5px;font-size:16px">
&#8681; Tai file Excel (.xlsx)</a></p>
<table border="1" cellpadding="6" style="border-collapse:collapse;width:100%;font-size:13px">
<tr style="background:#f0f0f0"><th>Thoi gian</th><th>Cau hoi</th><th>Cau tra loi</th></tr>
{rows}
</table>
<p style="color:gray">Chi hien 100 ban ghi gan nhat. Tai Excel de xem tat ca.</p>
</body></html>'''

@app.route('/logs/download')
def logs_download():
    pw = request.args.get('pw', '')
    if pw != LOG_PASSWORD:
        return 'Khong co quyen truy cap', 403
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = 'Log Cau Hoi Zalo'
        # Header
        headers = ['STT', 'Thoi gian', 'User ID', 'Cau hoi', 'Cau tra loi']
        ws.append(headers)
        for cell in ws[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill('solid', fgColor='1E6B3C')
            cell.alignment = Alignment(horizontal='center')
        # Du lieu
        for i, e in enumerate(_qa_log, 1):
            ws.append([i, e['time'], e['user_id'], e['question'], e['answer']])
        # Do rong cot
        ws.column_dimensions['A'].width = 6
        ws.column_dimensions['B'].width = 20
        ws.column_dimensions['C'].width = 20
        ws.column_dimensions['D'].width = 45
        ws.column_dimensions['E'].width = 60
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        fname = f"LogCauHoi_NhatGo_{_dt.datetime.now().strftime('%Y%m%d')}.xlsx"
        return Response(
            buf.read(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={'Content-Disposition': f'attachment; filename="{fname}"'}
        )
    except ImportError:
        return 'Chua cai openpyxl. Chay: pip install openpyxl', 500

# ── MAIN ─────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    print("=" * 62)
    print("  Tro Ly Nhat Go v5.0")
    print("  Streaming | Browser TTS | HTTPS | 2018-2026")
    print("=" * 62)
    print(f"  Folder 7 : {FOLDER7_PATH}")
    print(f"  Cache    : {INDEX_CACHE}")
    print(f"  Chunks   : {len(CHUNKS)}")

    PORT = int(os.environ.get('PORT', 5000))

    # Tren Render.com: bien RENDER=true -> chay HTTP (Render tu xu ly HTTPS)
    IS_CLOUD = bool(os.environ.get('RENDER', ''))

    if IS_CLOUD:
        print(f"  Cloud    : http://0.0.0.0:{PORT}")
        print("  HTTPS duoc xu ly boi Render.com")
        print("=" * 62)
        app.run(host='0.0.0.0', port=PORT, debug=False)
    else:
        ssl_ok = ensure_ssl()
        if ssl_ok:
            print("  HTTPS    : https://localhost:5000")
            print("  LAN      : https://<IP-may-tinh>:5000")
            print("  Lan dau  : Trinh duyet se canh bao cert tu ky -> Nhan 'Advanced > Proceed'")
            print("=" * 62)
            app.run(host='0.0.0.0', port=PORT, debug=False,
                    ssl_context=(SSL_CERT, SSL_KEY))
        else:
            print("  HTTP     : http://localhost:5000")
            print("  LAN      : http://<IP-may-tinh>:5000")
            print("  Luu y    : Micro chi hoat dong tren may chu (localhost)")
            print("=" * 62)
            app.run(host='0.0.0.0', port=PORT, debug=False)
